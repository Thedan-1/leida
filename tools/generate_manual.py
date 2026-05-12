import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

def create_placeholder_image(text, filename, size=(1000, 600), bg_color=(15, 23, 42)):
    img = Image.new('RGB', size, color=bg_color)
    draw = ImageDraw.Draw(img)
    # 模拟仪表盘的科技网格和线条
    for i in range(0, size[0], 50):
        draw.line([(i, 0), (i, size[1])], fill=(30, 41, 59), width=1)
    for j in range(0, size[1], 50):
        draw.line([(0, j), (size[0], j)], fill=(30, 41, 59), width=1)
    
    # 绘制科技感色块和占位标题
    draw.rectangle([50, 50, size[0]-50, size[1]-50], outline=(56, 189, 248), width=3)
    draw.line([(0,0), (50, 50)], fill=(56, 189, 248), width=2)
    draw.line([(size[0],0), (size[0]-50, 50)], fill=(56, 189, 248), width=2)
    
    try:
        font = ImageFont.truetype("Arial.ttf", 60)
    except:
        font = ImageFont.load_default()
        
    w, h = draw.textbbox((0, 0), text, font=font)[2:4]
    draw.text(((size[0]-w)/2, (size[1]-h)/2), text, fill=(56, 189, 248), font=font)
    img.save(filename)

def create_manual():
    doc = Document()
    
    # Cover Page
    title = doc.add_heading('工业元宇宙数字孪生SCADA系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('\n\n\n系统使用与操作手册 (V 3.0)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info = doc.add_paragraph('\n\n日期: 2026-04\n机密级别: 内部与客户专供')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    sections = [
        {
            "title": "一、 态势总览 (Overview - SCADA.CORE)",
            "desc": "本页面位于系统入口首屏，采用最新的工业级科技风格（Cyber-Industrial Theme）引擎进行实时数据渲染。提供了全天候的数据吞吐率诊断和全机组健康雷达评估。\n\n核心特性：\n1. TPS与报警数高亮挂载\n2. 传感器拟合参数模拟推演预测\n3. 健康雷达与时间帧日志分析\n4. 快速干预工单机动响应集群跟踪",
            "img_text": "【 态势总览 - 科幻大屏视图 】"
        },
        {
            "title": "二、 实时监测 (Realtime Monitoring)",
            "desc": "实时监测子系统下探至设备的最底级，为操作长和系统专家提供原生态时序数据流（Time-Series Data Stream）。\n\n核心特性：\n1. 具备毫秒级反应特征的管网压力和炉箱状态动态图表\n2. 纯数据密集型（Data-Dense）阵列模式排版，极大剥离光层修饰，保障数值精确展示\n3. 高亮闪烁报警和极限峰值追根溯源",
            "img_text": "【 实时监测 - 密集型时序视图 】"
        },
        {
            "title": "三、 控制中心 (Command Center)",
            "desc": "控制中心具有直接反向渗透设备的干预能力引擎。提供了关键致动器的启停和限值参数保护阀门调节枢纽。\n\n核心特性：\n1. 权限隔离架构：基于数字签章保护的命令下发机制\n2. 主阀门、高压水泵启停与备用切换指令集\n3. 带有硬锁阻断（Hardware Lock Prevent）的双重确认面板",
            "img_text": "【 控制中心 - 制动与调整枢纽 】"
        },
        {
            "title": "四、 演练预案 (Drill Center)",
            "desc": "用于应对最极端工业安全生产停摆事件的兵推中心。所有预案均与国家特种设备安全规范完全耦合。\n\n核心特性：\n1. 高峰期压力泄露模拟兵推与疏散逻辑图\n2. 人工介入时间倒数强制评估（Response Evaluation）\n3. 消防与防爆策略协同联动预案卡片",
            "img_text": "【 演练中心 - 安全预判模拟推演 】"
        },
        {
            "title": "五、 告警与工单中心 (Alarms & Workorders)",
            "desc": "以闭环管理（Closed-loop Management）为核心的生命周期系统。\n\n核心特性：\n1. 告警严重性分级（P0至P4）、未响应智能升级（Escalation）策略\n2. 人与机器协同的分配逻辑派单，支持动态接手与延期上报\n3. 处理时间MTTR与平均无故障时间MTBF双线挂载核算",
            "img_text": "【 日志工单 - 维保闭环执行墙 】"
        }
    ]

    for idx, sec in enumerate(sections):
        # Add heading
        heading = doc.add_heading(sec['title'], level=1)
        
        # Add description
        p = doc.add_paragraph(sec['desc'])
        p.style.font.size = Pt(12)
        
        # Create and add image
        img_temp = f'temp_img_{idx}.png'
        create_placeholder_image(sec['img_text'], img_temp)
        doc.add_picture(img_temp, width=Inches(6.0))
        
        # Add Page Break
        if idx < len(sections) - 1:
            doc.add_page_break()
            
        os.remove(img_temp)

    # Save documentation
    save_path = '/Users/thedan/Desktop/leida/系统使用说明_工业级大屏.docx'
    doc.save(save_path)
    print(f"Document Generated: {save_path}")

if __name__ == '__main__':
    create_manual()